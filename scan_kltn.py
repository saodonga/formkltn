"""
scan_kltn.py — Công cụ quét thư mục nộp bài KLTN/TTNN
======================================================
Đọc thông tin sinh viên từ:
  1. Tên file  (Họ tên, MSSV, Tên đề tài)
  2. Tên thư mục  (GVHD)
  3. Nội dung docx  (fallback nếu tên file không đủ thông tin)

Xuất ra file Excel có định dạng đẹp.

Cách dùng:
  python scan_kltn.py                        # Chọn thư mục qua hộp thoại
  python scan_kltn.py "F:\\KTS TTNN 2026"   # Truyền đường dẫn trực tiếp
"""

import os
import re
import sys
import unicodedata
from pathlib import Path
from datetime import datetime

try:
    import docx
except ImportError:
    print("[!] Thiếu thư viện python-docx. Đang cài...")
    os.system(f"{sys.executable} -m pip install python-docx -q")
    import docx

try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
except ImportError:
    print("[!] Thiếu thư viện openpyxl. Đang cài...")
    os.system(f"{sys.executable} -m pip install openpyxl -q")
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter


# ─────────────────────────────────────────────────────────────────
# 1. PARSE TÊN FILE
# ─────────────────────────────────────────────────────────────────

def parse_filename(filename: str) -> dict:
    """
    Phân tích tên file theo các pattern phổ biến:

    Pattern 1: 64KTS_HOÀNG THÙY LINH_2254105300_Báo cáo thực tập - Tên đề tài.docx
    Pattern 2: NGUYEN VAN A_2254100001_Ten de tai.docx
    Pattern 3: 2254100001_NGUYEN VAN A_ten de tai.docx
    Pattern 4: Chỉ tên file thông thường (fallback)
    """
    stem = Path(filename).stem  # bỏ .docx
    result = {
        'ho_ten': '',
        'mssv': '',
        'ten_de_tai_file': '',
        'raw_filename': stem,
    }

    # Tách theo dấu gạch dưới
    parts = [p.strip() for p in stem.split('_') if p.strip()]

    # Tìm MSSV: chuỗi 10 chữ số liên tiếp
    mssv_pattern = re.compile(r'\b\d{10}\b')
    # Tìm mã lớp: dạng 64KTS, 65KDTM, ... (chữ số + chữ cái)
    maloppart_pattern = re.compile(r'^\d{2}[A-Z]{2,6}$', re.IGNORECASE)

    mssv_idx = -1
    for i, part in enumerate(parts):
        m = mssv_pattern.search(part)
        if m:
            result['mssv'] = m.group()
            mssv_idx = i
            break

    # Xác định họ tên và đề tài dựa vào vị trí MSSV
    if mssv_idx >= 0:
        # Lấy phần trước MSSV, bỏ mã lớp nếu có
        before_mssv = parts[:mssv_idx]
        after_mssv  = parts[mssv_idx+1:]

        # Lọc bỏ mã lớp (vd: 64KTS)
        name_parts = [p for p in before_mssv if not maloppart_pattern.match(p)]
        result['ho_ten'] = ' '.join(name_parts).strip()

        # Phần sau MSSV là tên đề tài (có thể chứa dấu - để tách)
        de_tai_raw = ' '.join(after_mssv).strip()
        # Nếu có dạng "Báo cáo thực tập ngành - Tên đề tài", lấy phần sau dấu " - "
        # Nếu phần sau dấu " - " là tên người → dùng phần trước làm đề tài
        if ' - ' in de_tai_raw:
            de_tai_raw = de_tai_raw.split(' - ', 1)[1].strip()
        result['ten_de_tai_file'] = de_tai_raw

    else:
        # Không tìm thấy MSSV trong tên file
        # Thử tìm họ tên (chuỗi chữ hoa) và đề tài
        if len(parts) >= 2:
            # Giả sử phần đầu là họ tên nếu toàn chữ cái (không số)
            if re.match(r'^[A-ZÀÁẢÃẠĂẮẶẲẴẦẤẬẨẪÂĐÊẾỆỀỂỄÔỐỘỒỔỖƠỚỢỜỞỠƯỨỰỪỬỮÀÁẢÃẠĂẮẶẲẴẦẤẬẨẪÂĐÊẾỆỀỂỄÔỐỘỒỔỖƠỚỢỜỞỠƯỨỰỪỬỮ\s]+$', parts[0], re.UNICODE):
                result['ho_ten'] = parts[0]
                result['ten_de_tai_file'] = ' '.join(parts[1:])
            else:
                result['ten_de_tai_file'] = stem
        else:
            result['ten_de_tai_file'] = stem

    return result


# ─────────────────────────────────────────────────────────────────
# 2. PARSE TÊN THƯ MỤC → GVHD
# ─────────────────────────────────────────────────────────────────

def extract_gvhd_from_folder(folder_name: str) -> str:
    """
    Trích xuất tên GVHD từ tên thư mục.
    Ví dụ:
      'SV nhóm cô Dung'    → 'Cô Dung'
      'SV nhóm thầy Minh'  → 'Thầy Minh'
      'Nhóm GV Nguyễn Văn A' → 'Nguyễn Văn A'
      'cô Lan'             → 'Cô Lan'
    """
    name = folder_name.strip()

    # Pattern: "... cô/thầy/TS/PGS ..."
    patterns = [
        r'(?:nhóm\s+)?(?:của\s+)?((?:cô|thầy|TS\.?|PGS\.?TS\.?|GS\.?|ThS\.?)\s+[\w\sÀ-ỹ]+)',
        r'(?:SV\s+)?nhóm\s+([\w\sÀ-ỹ]+)',
        r'GV[:\s]+([\w\sÀ-ỹ]+)',
    ]
    for pat in patterns:
        m = re.search(pat, name, re.IGNORECASE | re.UNICODE)
        if m:
            gvhd = m.group(1).strip()
            # Capitalize đúng
            gvhd = ' '.join(w.capitalize() for w in gvhd.split())
            return gvhd

    # Fallback: trả về tên thư mục gốc, bỏ prefix "SV " nếu có
    name = re.sub(r'^SV\s+', '', name, flags=re.IGNORECASE).strip()
    return name


# ─────────────────────────────────────────────────────────────────
# 3. ĐỌC NỘI DUNG DOCX
# ─────────────────────────────────────────────────────────────────

def extract_from_docx(filepath: str) -> dict:
    """
    Đọc nội dung file docx để lấy thêm thông tin từ trang bìa.

    Cấu trúc trang bìa phổ biến ngành KTS:
        BÁO CÁO THỰC TẬP
        NGÀNH KINH TẾ SỐ
        <Tên đề tài — có thể nhiều dòng>
        ...
        Giảng viên hướng dẫn: ...
        Sinh viên thực hiện:   ...
        MSSV:                  ...
    """
    result = {
        'ho_ten_doc': '',
        'mssv_doc': '',
        'gvhd_doc': '',
        'ten_de_tai_doc': '',
        'error': '',
    }
    try:
        doc = docx.Document(filepath)

        # Thu thập toàn bộ paragraph text từ trang đầu (bìa)
        all_paras = []
        for para in doc.paragraphs[:80]:
            text = para.text.strip()
            all_paras.append(text)  # giữ cả dòng rỗng để biết khoảng cách

        # Lấy text từ bảng trong trang bìa (một số mẫu dùng bảng)
        table_lines = []
        for table in doc.tables[:5]:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text = para.text.strip()
                        if text:
                            table_lines.append(text)

        # Kết hợp
        lines_nonempty = [l for l in all_paras if l]
        full_text = '\n'.join(lines_nonempty[:50])
        full_text_with_tables = full_text + '\n' + '\n'.join(table_lines)

        # ── 1. TÊN ĐỀ TÀI ────────────────────────────────────────────
        # Chiến lược: Tìm anchor "BÁO CÁO THỰC TẬP" hoặc "NGÀNH KINH TẾ SỐ"
        # rồi lấy các dòng liền sau (bỏ qua dòng rỗng) cho đến khi gặp
        # dòng chứa từ khóa tiếp theo (GVHD, Sinh viên, Lớp, v.v.)
        ANCHOR_PATTERNS = [
            r'BÁO CÁO THỰC TẬP',
            r'BÁO CÁO THỰC TẬP NGÀNH',
            r'NGÀNH KINH TẾ SỐ',
            r'BÁO CÁO TỐT NGHIỆP',
            r'KHÓA LUẬN TỐT NGHIỆP',
        ]
        STOP_PATTERNS = [
            r'(?:Giảng viên|GVHD|GV\.?HD|Hướng dẫn)',
            r'(?:Sinh viên|Họ và tên|Họ tên|SV thực hiện)',
            r'(?:Lớp|MSSV|Mã số|Khóa|Năm học)',
            r'(?:TP\.|Hà Nội|TP HCM|Đà Nẵng|Cần Thơ).*\d{4}',
            r'\d{4}',  # năm riêng lẻ
        ]

        ten_de_tai = ''
        # Duyệt qua từng dòng để tìm anchor
        for anchor_pat in ANCHOR_PATTERNS:
            anchor_idx = None
            for i, line in enumerate(lines_nonempty):
                if re.search(anchor_pat, line, re.IGNORECASE):
                    anchor_idx = i
                    break
            if anchor_idx is None:
                continue

            # Tìm cuối khối anchor (có thể 1-2 dòng liên tiếp như "BÁO CÁO" + "THỰC TẬP NGÀNH KTS")
            end_anchor = anchor_idx
            for j in range(anchor_idx, min(anchor_idx + 3, len(lines_nonempty))):
                if any(re.search(ap, lines_nonempty[j], re.IGNORECASE) for ap in ANCHOR_PATTERNS):
                    end_anchor = j

            # Lấy các dòng sau anchor cho đến STOP
            title_lines = []
            for line in lines_nonempty[end_anchor + 1:end_anchor + 10]:
                if any(re.search(sp, line, re.IGNORECASE) for sp in STOP_PATTERNS):
                    break
                if len(line) >= 5:  # bỏ dòng quá ngắn
                    title_lines.append(line)

            if title_lines:
                ten_de_tai = ' '.join(title_lines).strip()
                break

        # Fallback: tìm theo keyword "Đề tài:", "Tên đề tài:"
        if not ten_de_tai:
            m = re.search(
                r'(?:Đề tài|Tên đề tài|ĐỀ TÀI)[:\s]+([^\n]{10,200})',
                full_text_with_tables, re.IGNORECASE
            )
            if m:
                ten_de_tai = m.group(1).strip()

        # Bỏ dấu nháy đơn/đôi bao quanh hoặc lẫn trong tên đề tài
        ten_de_tai = ten_de_tai.replace('"', '').replace('"', '').replace('"', '')
        ten_de_tai = ten_de_tai.replace("'", '').replace('\u2018', '').replace('\u2019', '')
        ten_de_tai = ten_de_tai.strip(' \t\n\r.,;:')
        result['ten_de_tai_doc'] = ten_de_tai[:250]

        # ── 2. GVHD ──────────────────────────────────────────────────
        gvhd_patterns = [
            r'(?:Giảng viên hướng dẫn|GVHD|GV\.?HD|Giáo viên hướng dẫn|Người hướng dẫn)[:\s.]+([^\n]{5,80})',
            r'(?:Hướng dẫn khoa học)[:\s]+([^\n]{5,80})',
            r'(?:Hướng dẫn)[:\s]+([^\n]{5,80})',
        ]
        for pat in gvhd_patterns:
            m = re.search(pat, full_text_with_tables, re.IGNORECASE)
            if m:
                gvhd = m.group(1).strip()
                # Bỏ học hàm/học vị đứng đầu (TS., ThS., PGS.TS., ...)
                gvhd = re.sub(r'^(PGS\.?TS\.?|GS\.?TS\.?|TS\.?|ThS\.?|GS\.?)\s+', '', gvhd, flags=re.IGNORECASE).strip()
                # Bỏ nếu chỉ chứa dấu chấm hoặc quá ngắn
                if len(gvhd) >= 3:
                    result['gvhd_doc'] = gvhd[:60]
                    break

        # ── 3. HỌ TÊN SINH VIÊN ──────────────────────────────────────
        ten_patterns = [
            r'(?:Sinh viên thực hiện|Sinh viên|Họ và tên SV|Họ và tên|Tên sinh viên)[:\s]+([^\n]{5,60})',
            r'(?:Người thực hiện|Học viên)[:\s]+([^\n]{5,60})',
        ]
        for pat in ten_patterns:
            m = re.search(pat, full_text_with_tables, re.IGNORECASE)
            if m:
                ho_ten = m.group(1).strip()
                # Bỏ nếu chứa số (có thể bị dính MSSV)
                ho_ten = re.split(r'\s{2,}|\t|\d{5,}', ho_ten)[0].strip()
                if len(ho_ten) >= 3:
                    result['ho_ten_doc'] = ho_ten[:60]
                    break

        # ── 4. MSSV ──────────────────────────────────────────────────
        # Ưu tiên tìm sau nhãn "MSSV:"
        mssv_labeled = re.search(
            r'(?:MSSV|Mã số SV|Mã sinh viên)[:\s]+([\d]{8,12})',
            full_text_with_tables, re.IGNORECASE
        )
        if mssv_labeled:
            result['mssv_doc'] = mssv_labeled.group(1).strip()
        else:
            # Fallback: tìm chuỗi 10 số liên tiếp bất kỳ
            m = re.search(r'\b(\d{10})\b', full_text_with_tables)
            if m:
                result['mssv_doc'] = m.group(1)

    except Exception as e:
        result['error'] = str(e)[:150]
    return result


# ─────────────────────────────────────────────────────────────────
# 4. TRÍCH TÊN CÔNG TY TỪ TÊN ĐỀ TÀI
# ─────────────────────────────────────────────────────────────────

def extract_ten_cong_ty(ten_de_tai: str) -> str:
    """
    Tách tên công ty / tổ chức từ tên đề tài.
    Tìm từ khoá bắt đầu (công ty, doanh nghiệp, ngân hàng, ...)
    rồi lấy toàn bộ phần còn lại.

    Ví dụ:
      "Giới thiệu mô hình ... của công ty cổ phần Global Aspire"
          → "công ty cổ phần Global Aspire"
      "Hoàn thiện chiến lược tại Công ty TNHH ABC Việt Nam"
          → "Công ty TNHH ABC Việt Nam"
      "Phân tích hoạt động tại Ngân hàng Vietcombank – Chi nhánh HCM"
          → "Ngân hàng Vietcombank – Chi nhánh HCM"
    """
    if not ten_de_tai:
        return ''

    # Các từ khoá mở đầu tên tổ chức (theo thứ tự ưu tiên)
    keywords = [
        r'công ty cổ phần',
        r'công ty tnhh',
        r'công ty hợp danh',
        r'công ty tư nhân',
        r'công ty',
        r'doanh nghiệp',
        r'ngân hàng',
        r'tập đoàn',
        r'tổng công ty',
        r'chi nhánh',
        r'siêu thị',
        r'chuỗi',
        r'trung tâm',
        r'cửa hàng',
    ]

    text_lower = ten_de_tai.lower()
    best_idx = len(ten_de_tai)  # mặc định không tìm thấy

    for kw in keywords:
        m = re.search(kw, text_lower)
        if m and m.start() < best_idx:
            best_idx = m.start()

    if best_idx < len(ten_de_tai):
        ten_ct = ten_de_tai[best_idx:].strip(' \t\n.,;:')
        # Bỏ dấu ngoặc và ký tự thừa ở cuối
        ten_ct = re.sub(r'[\(\)\[\]]+$', '', ten_ct).strip()
        return ten_ct

    return ''




def scan_folder(root_path: str, read_docx: bool = True) -> list[dict]:
    """
    Quét đệ quy các thư mục con, thu thập thông tin từ file .docx
    """
    root = Path(root_path)
    records = []
    total_files = 0
    error_files = 0

    print(f"\n{'='*60}")
    print(f"📂 Thư mục gốc: {root}")
    print(f"{'='*60}\n")

    # Liệt kê các thư mục con trực tiếp (nhóm GVHD)
    subfolders = [f for f in root.iterdir() if f.is_dir()]
    if not subfolders:
        # Không có thư mục con → quét thẳng root
        subfolders = [root]

    print(f"🗂️  Tìm thấy {len(subfolders)} thư mục:\n")
    for sf in sorted(subfolders):
        gvhd_folder = extract_gvhd_from_folder(sf.name)
        docx_files = list(sf.glob('**/*.docx'))
        # Bỏ qua file ~$ (file tạm của Word đang mở)
        docx_files = [f for f in docx_files if not f.name.startswith('~$')]
        print(f"  📁 {sf.name}")
        print(f"      → GVHD (từ tên thư mục): {gvhd_folder}")
        print(f"      → Số file .docx: {len(docx_files)}\n")

        for filepath in sorted(docx_files):
            total_files += 1
            # Parse từ tên file
            file_info = parse_filename(filepath.name)
            # Đọc nội dung docx
            doc_info = {}
            if read_docx:
                print(f"    📄 Đọc: {filepath.name[:70]}...")
                doc_info = extract_from_docx(str(filepath))
                if doc_info.get('error'):
                    print(f"       ⚠️  Lỗi: {doc_info['error']}")
                    error_files += 1

            # Kết hợp thông tin — ưu tiên: docx > tên file > tên thư mục
            ho_ten = (doc_info.get('ho_ten_doc') or file_info['ho_ten']).strip()
            mssv   = (doc_info.get('mssv_doc')   or file_info['mssv']).strip()
            gvhd   = (doc_info.get('gvhd_doc')   or gvhd_folder).strip()
            ten_de_tai = (doc_info.get('ten_de_tai_doc') or file_info['ten_de_tai_file']).strip()
            ten_cong_ty = extract_ten_cong_ty(ten_de_tai)

            records.append({
                'STT': total_files,
                'Thư mục nhóm': sf.name,
                'GVHD': gvhd,
                'Họ và tên SV': ho_ten,
                'MSSV': mssv,
                'Tên đề tài': ten_de_tai,
                'Tên công ty': ten_cong_ty,
                # Nguồn gốc để kiểm tra
                'Tên file gốc': filepath.name,
                'HoTen_TuFile': file_info['ho_ten'],
                'MSSV_TuFile': file_info['mssv'],
                'DeTai_TuFile': file_info['ten_de_tai_file'],
                'HoTen_TuDoc': doc_info.get('ho_ten_doc', ''),
                'MSSV_TuDoc': doc_info.get('mssv_doc', ''),
                'GVHD_TuDoc': doc_info.get('gvhd_doc', ''),
                'DeTai_TuDoc': doc_info.get('ten_de_tai_doc', ''),
                'Loi_DocFile': doc_info.get('error', ''),
            })

    print(f"\n{'='*60}")
    print(f"✅ Tổng: {total_files} file | Lỗi đọc: {error_files}")
    print(f"{'='*60}\n")
    return records


# ─────────────────────────────────────────────────────────────────
# 5. XUẤT EXCEL
# ─────────────────────────────────────────────────────────────────

def export_excel(records: list[dict], output_path: str):
    wb = Workbook()

    thin = Side(style='thin', color='BFBFBF')
    bdr  = Border(left=thin, right=thin, top=thin, bottom=thin)

    def hcell(ws, r, c, val, fill='2C3E50', color='FFFFFF', size=11, bold=True, center=True):
        cell = ws.cell(row=r, column=c, value=val)
        cell.font = Font(name='Arial', size=size, bold=bold, color=color)
        cell.fill = PatternFill('solid', start_color=fill)
        cell.alignment = Alignment(horizontal='center' if center else 'left',
                                   vertical='center', wrap_text=True)
        cell.border = bdr
        return cell

    def dcell(ws, r, c, val, fill='FFFFFF', color='000000', size=11, bold=False,
              center=False, italic=False):
        cell = ws.cell(row=r, column=c, value=val)
        cell.font = Font(name='Arial', size=size, bold=bold, color=color, italic=italic)
        cell.fill = PatternFill('solid', start_color=fill)
        cell.alignment = Alignment(horizontal='center' if center else 'left',
                                   vertical='center', wrap_text=True)
        cell.border = bdr
        return cell

    # ── Sheet 1: Danh sách chính ──────────────────────────────────
    ws1 = wb.active
    ws1.title = "Danh sách SV"

    ws1.merge_cells('A1:G1')
    c = ws1['A1']
    c.value = "DANH SÁCH SINH VIÊN NỘP BÁO CÁO THỰC TẬP NGÀNH — NGÀNH KINH TẾ SỐ"
    c.font = Font(name='Arial', size=13, bold=True, color='FFFFFF')
    c.fill = PatternFill('solid', start_color='1A2A3A')
    c.alignment = Alignment(horizontal='center', vertical='center')
    ws1.row_dimensions[1].height = 34

    ws1.merge_cells('A2:G2')
    ts = datetime.now().strftime('%d/%m/%Y %H:%M')
    c = ws1['A2']
    c.value = f"Xuất lúc: {ts}  ·  Tổng số: {len(records)} sinh viên"
    c.font = Font(name='Arial', size=10, italic=True, color='777777')
    c.fill = PatternFill('solid', start_color='F8F9FA')
    c.alignment = Alignment(horizontal='center', vertical='center')
    ws1.row_dimensions[2].height = 20

    headers = ['STT', 'GVHD', 'Họ và tên SV', 'MSSV', 'Tên đề tài', 'Tên công ty', 'Thư mục nhóm']
    ws1.row_dimensions[3].height = 30
    for ci, h in enumerate(headers, 1):
        hcell(ws1, 3, ci, h)

    # Group colors by GVHD
    gvhd_list = list(dict.fromkeys(r['GVHD'] for r in records))
    palette = ['EBF5FB', 'D9EDD6', 'FDECEA', 'F3E8FB', 'FEF9E7', 'EAF4FB',
               'FFF3E8', 'E8F5E9', 'FCE4EC', 'E0F7FA']
    gvhd_color = {g: palette[i % len(palette)] for i, g in enumerate(gvhd_list)}

    for i, rec in enumerate(records):
        r = i + 4
        ws1.row_dimensions[r].height = 20
        fill = gvhd_color.get(rec['GVHD'], 'FFFFFF')
        alt  = 'F8F8F8' if i % 2 == 0 else 'FFFFFF'

        dcell(ws1, r, 1, rec['STT'],            fill=alt,      center=True, bold=True)
        dcell(ws1, r, 2, rec['GVHD'],           fill=fill,     bold=True, size=11)
        dcell(ws1, r, 3, rec['Họ và tên SV'],   fill=alt,      size=11)
        dcell(ws1, r, 4, rec['MSSV'],           fill=alt,      center=True, size=11)
        dcell(ws1, r, 5, rec['Tên đề tài'],     fill=alt,      size=11)
        dcell(ws1, r, 6, rec['Tên công ty'],    fill='FFFDE7', size=11, italic=True, color='7E5109')
        dcell(ws1, r, 7, rec['Thư mục nhóm'],  fill='F0F0F0', size=10, italic=True, color='777777')

    ws1.column_dimensions['A'].width = 5
    ws1.column_dimensions['B'].width = 22
    ws1.column_dimensions['C'].width = 22
    ws1.column_dimensions['D'].width = 13
    ws1.column_dimensions['E'].width = 50
    ws1.column_dimensions['F'].width = 30
    ws1.column_dimensions['G'].width = 22
    ws1.auto_filter.ref = 'A3:G3'
    ws1.freeze_panes   = 'A4'
    ws1.sheet_properties.tabColor = '1F4E79'

    # ── Sheet 2: Dữ liệu đầy đủ (để kiểm tra/debug) ──────────────
    ws2 = wb.create_sheet("Dữ liệu chi tiết")

    ws2.merge_cells('A1:N1')
    c = ws2['A1']
    c.value = "DỮ LIỆU CHI TIẾT — NGUỒN GỐC TỪNG TRƯỜNG (để kiểm tra và hiệu chỉnh)"
    c.font = Font(name='Arial', size=12, bold=True, color='FFFFFF')
    c.fill = PatternFill('solid', start_color='1A2A3A')
    c.alignment = Alignment(horizontal='center', vertical='center')
    ws2.row_dimensions[1].height = 30

    hdrs2 = ['STT', 'Thư mục nhóm', 'GVHD (kết hợp)',
             'Họ tên (kết hợp)', 'MSSV (kết hợp)', 'Đề tài (kết hợp)',
             'HoTen từ File', 'MSSV từ File', 'ĐT từ File',
             'HoTen từ Doc', 'MSSV từ Doc', 'GVHD từ Doc', 'ĐT từ Doc',
             'Tên file gốc', 'Lỗi']
    ws2.row_dimensions[2].height = 28
    for ci, h in enumerate(hdrs2, 1):
        hcell(ws2, 2, ci, h, size=10)

    keys = ['STT', 'Thư mục nhóm', 'GVHD', 'Họ và tên SV', 'MSSV', 'Tên đề tài',
            'HoTen_TuFile', 'MSSV_TuFile', 'DeTai_TuFile',
            'HoTen_TuDoc', 'MSSV_TuDoc', 'GVHD_TuDoc', 'DeTai_TuDoc',
            'Tên file gốc', 'Loi_DocFile']

    for i, rec in enumerate(records):
        r = i + 3
        ws2.row_dimensions[r].height = 18
        fill = 'FAFAFA' if i % 2 == 0 else 'FFFFFF'
        for ci, key in enumerate(keys, 1):
            val = rec.get(key, '')
            c = ws2.cell(row=r, column=ci, value=val)
            c.font = Font(name='Arial', size=10,
                          color='CC0000' if key == 'Loi_DocFile' and val else '000000')
            c.fill = PatternFill('solid', start_color=fill)
            c.alignment = Alignment(vertical='center', wrap_text=False)
            c.border = bdr

    for ci, w in enumerate([4,22,22,22,13,40,20,13,40,20,13,22,40,40,20], 1):
        ws2.column_dimensions[get_column_letter(ci)].width = w
    ws2.freeze_panes = 'C3'
    ws2.sheet_properties.tabColor = '1E6B3A'

    # ── Sheet 3: Thống kê theo GVHD ──────────────────────────────
    ws3 = wb.create_sheet("Thống kê GVHD")

    ws3.merge_cells('A1:C1')
    c = ws3['A1']
    c.value = "THỐNG KÊ SỐ LƯỢNG SINH VIÊN THEO GVHD"
    c.font = Font(name='Arial', size=12, bold=True, color='FFFFFF')
    c.fill = PatternFill('solid', start_color='1A2A3A')
    c.alignment = Alignment(horizontal='center', vertical='center')
    ws3.row_dimensions[1].height = 30

    for ci, h in enumerate(['GVHD', 'Số lượng SV', 'Tên thư mục nhóm'], 1):
        hcell(ws3, 2, ci, h)
    ws3.row_dimensions[2].height = 26

    from collections import Counter
    gvhd_counts = Counter(r['GVHD'] for r in records)
    gvhd_folders = {}
    for r in records:
        gvhd_folders.setdefault(r['GVHD'], set()).add(r['Thư mục nhóm'])

    for ri, (gvhd, cnt) in enumerate(sorted(gvhd_counts.items())):
        r = ri + 3
        ws3.row_dimensions[r].height = 24
        fill = palette[ri % len(palette)]
        c = ws3.cell(row=r, column=1, value=gvhd)
        c.font = Font(name='Arial', size=11, bold=True); c.fill = PatternFill('solid', start_color=fill)
        c.alignment = Alignment(vertical='center'); c.border = bdr
        c = ws3.cell(row=r, column=2, value=cnt)
        c.font = Font(name='Arial', size=12, bold=True, color='1F4E79')
        c.fill = PatternFill('solid', start_color=fill)
        c.alignment = Alignment(horizontal='center', vertical='center'); c.border = bdr
        c = ws3.cell(row=r, column=3, value=', '.join(sorted(gvhd_folders[gvhd])))
        c.font = Font(name='Arial', size=10, italic=True, color='555555')
        c.fill = PatternFill('solid', start_color='FAFAFA')
        c.alignment = Alignment(vertical='center', wrap_text=True); c.border = bdr

    total_r = len(gvhd_counts) + 3
    ws3.row_dimensions[total_r].height = 28
    for ci, val in enumerate(['TỔNG CỘNG', len(records), ''], 1):
        c = ws3.cell(row=total_r, column=ci, value=val)
        c.font = Font(name='Arial', size=12, bold=True, color='FFFFFF')
        c.fill = PatternFill('solid', start_color='2C3E50')
        c.alignment = Alignment(horizontal='center', vertical='center'); c.border = bdr

    ws3.column_dimensions['A'].width = 28
    ws3.column_dimensions['B'].width = 16
    ws3.column_dimensions['C'].width = 36
    ws3.sheet_properties.tabColor = '7B2C2C'

    wb.save(output_path)
    print(f"✅ Đã lưu Excel: {output_path}")


# ─────────────────────────────────────────────────────────────────
# 6. MAIN
# ─────────────────────────────────────────────────────────────────

def main():
    print("\n" + "="*60)
    print("  SCAN KLTN/TTNN — Quét dữ liệu sinh viên từ file .docx")
    print("="*60)

    # Lấy đường dẫn từ argument hoặc hỏi người dùng
    if len(sys.argv) > 1:
        root_path = sys.argv[1].strip('"').strip("'")
    else:
        # Thử dùng hộp thoại chọn thư mục (Windows/Mac/Linux)
        try:
            import tkinter as tk
            from tkinter import filedialog
            root_tk = tk.Tk()
            root_tk.withdraw()
            root_path = filedialog.askdirectory(
                title='Chọn thư mục gốc chứa các nhóm KLTN/TTNN'
            )
            root_tk.destroy()
            if not root_path:
                print("❌ Không chọn thư mục. Thoát.")
                sys.exit(0)
        except Exception:
            root_path = input("Nhập đường dẫn thư mục gốc: ").strip().strip('"')

    if not os.path.isdir(root_path):
        print(f"❌ Đường dẫn không tồn tại: {root_path}")
        sys.exit(1)

    print(f"\n⚙️  Bắt đầu quét: {root_path}")
    print("   (Quá trình đọc từng file docx có thể mất vài phút...)\n")

    records = scan_folder(root_path, read_docx=True)

    if not records:
        print("⚠️  Không tìm thấy file .docx nào!")
        sys.exit(0)

    # Đường dẫn file output
    output_path = os.path.join(
        root_path,
        f"TONG_HOP_SV_TTNN_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx"
    )

    export_excel(records, output_path)
    print(f"\n🎉 Hoàn tất! Mở file Excel tại:\n   {output_path}\n")

    # Tự động mở file Excel (Windows)
    if sys.platform == 'win32':
        try:
            os.startfile(output_path)
        except Exception:
            pass


if __name__ == '__main__':
    main()
